# Data Manipulation Libraries
import pandas as pd  # Data manipulation and analysis
import numpy as np  # Numerical operations on large multi-dimensional arrays and matrices

# Visualization Libraries
import matplotlib.pyplot as plt  # Plotting library for creating static, animated, and interactive visualizations
import seaborn as sns  # Statistical data visualization

# Statistical Analysis Libraries
from scipy.cluster.hierarchy import dendrogram, linkage  # Hierarchical clustering utilities
import scipy.cluster.hierarchy as sch  # Additional hierarchical clustering methods from SciPy
from scipy.stats import shapiro, chi2_contingency, mannwhitneyu  # Statistical tests

# Word Document Manipulation
from docx import Document  # For .docx file creation and manipulation
import docx.shared  # For resizing images in the Word document

# Date and Time Utilities
from datetime import datetime  # For working with date and time

# Machine Learning - Preprocessing and Model Evaluation
from sklearn.preprocessing import StandardScaler, LabelEncoder  # Standardization of features and label encoding
from sklearn.feature_selection import VarianceThreshold, mutual_info_classif  # Feature selection
from sklearn.model_selection import train_test_split  # For splitting the dataset into training and testing sets
from sklearn.linear_model import LogisticRegression, LinearRegression  # Logistic and Linear regression models for analysis
from sklearn.metrics import confusion_matrix, classification_report, accuracy_score, r2_score  # Evaluation metrics for classification and regression
from sklearn.ensemble import RandomForestClassifier  # Ensemble method for classification
from sklearn.cluster import KMeans, AgglomerativeClustering  # Clustering algorithms
from sklearn.decomposition import PCA  # Principal Component Analysis for dimensionality reduction
from sklearn.metrics import silhouette_score, davies_bouldin_score, calinski_harabasz_score  # Clustering evaluation metrics

# Data Serialization and File Handling
import os  # Provides a way of using operating system dependent functionality like reading or writing to the file system
import shutil  # High-level file operations like copying and removal


# Custom Functions for Document Generation and Display
def print_to_doc(doc, *args):
    """Prints given text to a docx document."""
    text = ' '.join(map(str, args))
    doc.add_paragraph(text)

def custom_display(doc, output_to_docx, *args):
    """Custom display function for printing or saving output to a docx document."""
    if output_to_docx == 1:
        for arg in args:
            if isinstance(arg, pd.DataFrame):
                buf = StringIO()
                arg.to_string(buf)
                text = buf.getvalue()
            else:
                text = str(arg)
            doc.add_paragraph(text)
    else:
        display(*args)

def save_document(doc, filename='output.docx'):
    """Saves the docx document with the current date and time."""
    current_date = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    doc.add_paragraph(f"Document saved on {current_date}")
    doc.save(filename)
    print(f"Document saved as {filename}")

def add_image_to_doc(doc, image_path):
    """Adds an image to the docx document."""
    doc.add_picture(image_path, width=docx.shared.Inches(6))
    doc.add_paragraph(f'Image: {os.path.basename(image_path)}')

# Descriptive Statistics
def descriptive_statistics(dataset, year, all_lists, doc=None):
    """Computes and displays descriptive statistics for a dataset."""
    columns_of_interest = ['Total_AP', 'Total_PP', 'Total_SPA', 'FinalGrade']
    categorical_attributes = all_lists[year].get('basic_lists').get('listCategorical')
    selected_columns = columns_of_interest + categorical_attributes

    filtered_dataset = dataset[selected_columns]

    numerical_stats = filtered_dataset[columns_of_interest].describe().transpose()
    numerical_stats['skew'] = filtered_dataset[columns_of_interest].skew()
    numerical_stats['kurt'] = filtered_dataset[columns_of_interest].kurt()

    categorical_stats = filtered_dataset[categorical_attributes].describe().transpose()

    descriptive_stats = pd.concat([numerical_stats, categorical_stats], axis=0)

    if doc:
        print_to_doc(doc, f"Descriptive statistics for year {year}:\n", descriptive_stats)
    else:
        print(f"Descriptive statistics for year {year}:\n", descriptive_stats)

    return descriptive_stats

# Function to print Pairplots between FinalGrade and the Total of Assessment Points
def multivariate_analysis(dataset, year, save_path='Charts/MultivariateAnalysis'):
    """Generates and saves a pairplot for selected numerical columns in the dataset."""
    if not os.path.exists(save_path):
        os.makedirs(save_path)

    selected_cols = ['FinalGrade', 'Total_AP', 'Total_PP', 'Total_SPA', 'Total_Bon']
    pairplot = sns.pairplot(dataset[selected_cols])
    pairplot.fig.suptitle(f'Pairplot for {year}', y=1.02)
    pairplot.fig.savefig(f'{save_path}/Pairplot_{year}.png', dpi=300)
    plt.show()

# Function to prepare data for Chi-squared test and check sufficiency
def prepare_chi_squared_data(df, column1, column2):
    contingency_table = pd.crosstab(df[column1], df[column2])
    if (contingency_table.sum(axis=1) < 5).any() or (contingency_table.sum(axis=0) < 5).any():
        return None  # Insufficient data for Chi-squared test
    return contingency_table

# Function to perform Chi-squared test for two categorical variables
def chi_squared_test(dataset, cat_var1, cat_var2, doc=None):
    """Performs and displays the result of a Chi-squared test between two categorical variables."""
    if cat_var1 in dataset.columns and cat_var2 in dataset.columns:
        contingency_table = prepare_chi_squared_data(dataset, cat_var1, cat_var2)
        if contingency_table is not None:
            chi2_stat, p_val, dof, _ = chi2_contingency(contingency_table)
            result = (
                f"Chi-squared test between {cat_var1} and {cat_var2}:\n"
                f"Chi-squared statistic: {chi2_stat}\n"
                f"Degrees of freedom: {dof}\n"
                f"P-value: {p_val}\n"
            )
            if doc:
                print_to_doc(doc, result)
            else:
                print(result)
        else:
            error_msg = f"Insufficient data for Chi-squared test between {cat_var1} and {cat_var2}."
            if doc:
                print_to_doc(doc, error_msg)
            else:
                print(error_msg)
    else:
        error_msg = f"One or both variables {cat_var1} and {cat_var2} are not in the dataset."
        if doc:
            print_to_doc(doc, error_msg)
        else:
            print(error_msg)

# Function to perform Chi-squared test for a categorical variable and binned numerical variable
def chi_squared_test_binned(dataset, cat_var, num_var, bins, doc=None):
    """Performs and displays the result of a Chi-squared test between a categorical variable and a binned numerical variable."""
    if cat_var in dataset.columns and num_var in dataset.columns:
        dataset['binned'] = pd.cut(dataset[num_var], bins)
        contingency_table = prepare_chi_squared_data(dataset, cat_var, 'binned')
        if contingency_table is not None:
            chi2_stat, p_val, dof, _ = chi2_contingency(contingency_table)
            result = (
                f"Chi-squared test between {cat_var} and {num_var} (binned):\n"
                f"Chi-squared statistic: {chi2_stat}\n"
                f"Degrees of freedom: {dof}\n"
                f"P-value: {p_val}\n"
            )
            if doc:
                print_to_doc(doc, result)
            else:
                print(result)
        else:
            error_msg = f"Insufficient data for Chi-squared test between {cat_var} and {num_var} (binned)."
            if doc:
                print_to_doc(doc, error_msg)
            else:
                print(error_msg)
    else:
        error_msg = f"One or both variables {cat_var} and {num_var} are not in the dataset."
        if doc:
            print_to_doc(doc, error_msg)
        else:
            print(error_msg)

# Function to calculate the number of bins for a numerical variable
def calculate_bins(dataset, num_var):
    """Calculates an appropriate number of bins for a numerical variable using the Freedman-Diaconis rule."""
    if num_var in dataset.columns:
        iqr = np.percentile(dataset[num_var], 75) - np.percentile(dataset[num_var], 25)
        bin_width = 2 * iqr / (len(dataset[num_var]) ** (1/3))
        data_range = dataset[num_var].max() - dataset[num_var].min()
        num_bins = int(np.ceil(data_range / bin_width))
        return num_bins
    else:
        return 10

# Function to perform Mann-Whitney U test between a categorical and a numerical variable
def mann_whitney_u_test(dataset, cat_var, num_var, doc=None):
    """Performs and displays the result of a Mann-Whitney U test between a categorical variable and a numerical variable."""
    if cat_var in dataset.columns and num_var in dataset.columns:
        group1 = dataset[dataset[cat_var] == 0][num_var]
        group2 = dataset[dataset[cat_var] == 1][num_var]
        if len(group1) > 0 and len(group2) > 0:
            u_statistic, p_value = mannwhitneyu(group1, group2)
            result = (
                f"Mann-Whitney U test results between {cat_var} and {num_var}:\n"
                f"U-statistic: {u_statistic}\n"
                f"P-value: {p_value}\n"
            )
            if doc:
                print_to_doc(doc, result)
            else:
                print(result)
        else:
            error_msg = f"At least one group has no data for {cat_var} and {num_var}. Mann-Whitney U test cannot be performed."
            if doc:
                print_to_doc(doc, error_msg)
            else:
                print(error_msg)
    else:
        error_msg = f"One or both variables {cat_var} and {num_var} are not in the dataset."
        if doc:
            print_to_doc(doc, error_msg)
        else:
            print(error_msg)

# Function to perform Chi-squared tests for pairs of categorical variables
def perform_chi_squared_tests_cat_pairs(dataset, cat_pairs, doc=None):
    """Performs Chi-squared tests for all pairs of categorical variables."""
    for cat_var1, cat_var2 in cat_pairs:
        chi_squared_test(dataset, cat_var1, cat_var2, doc)

# Function to perform Chi-squared tests for categorical vs. binned numerical variables
def perform_chi_squared_tests(dataset, cat_vars, num_vars, doc=None):
    """Performs Chi-squared tests for all pairs of categorical variables and binned numerical variables."""
    for cat_var in cat_vars:
        if cat_var in dataset.columns:
            for num_var in num_vars:
                if num_var in dataset.columns:
                    num_bins = calculate_bins(dataset, num_var)
                    chi_squared_test_binned(dataset, cat_var, num_var, num_bins, doc)

# Function to perform Mann-Whitney U tests for categorical vs. numerical variables
def perform_mann_whitney_u_tests(dataset, cat_vars, num_vars, doc=None):
    """Performs Mann-Whitney U tests for all pairs of categorical variables and numerical variables."""
    for cat_var in cat_vars:
        if cat_var in dataset.columns:
            for num_var in num_vars:
                if num_var in dataset.columns:
                    mann_whitney_u_test(dataset, cat_var, num_var, doc)

# Function to run all specified tests for a list of datasets
def run_all_tests(datasets, categorical_vars, numerical_vars, categorical_pairs, doc=None):
    """Runs all specified tests for each dataset in the datasets dictionary."""
    for year, dataset in datasets.items():
        header = f"\nYear: {year}\n{'-'*50}"
        if doc:
            print_to_doc(doc, header)
        else:
            print(header)

        if doc:
            print_to_doc(doc, "Chi-Squared Tests for Categorical Pairs:")
        perform_chi_squared_tests_cat_pairs(dataset, categorical_pairs, doc)

        if doc:
            print_to_doc(doc, "Chi-Squared Tests for Categorical vs Binned Numerical Variables:")
        perform_chi_squared_tests(dataset, categorical_vars, numerical_vars, doc)

        if doc:
            print_to_doc(doc, "Mann-Whitney U Tests for Categorical vs Numerical Variables:")
        perform_mann_whitney_u_tests(dataset, categorical_vars, numerical_vars, doc)

# Normality Check
def check_normality(dataset, num_columns, doc=None):
    """Checks for normality of numerical columns using the Shapiro-Wilk test."""
    normality_results = {}
    for col in num_columns:
        if col in dataset.columns:
            data = dataset[col].dropna()
            if data.nunique() > 1:
                stat, p_val = shapiro(data)
                normality_results[col] = {'W': stat, 'p-value': p_val}
            else:
                normality_results[col] = {'W': None, 'p-value': None}
    return normality_results

# Course and Gender ANOVA
def course_gender_anova(datasets, doc=None):
    """Performs ANOVA to study the effect of course and gender on the final grade."""
    for year, dataset in datasets.items():
        try:
            header = "----------------------------------------------------------------------------------"
            if doc:
                print_to_doc(doc, header)
            else:
                print(header)
            if doc:
                print_to_doc(doc, year)
            else:
                print(year)
            model = ols('FinalGrade ~ C(Course) + C(Sex) + C(Course):C(Sex)', data=dataset).fit()
            anova_table = sm.stats.anova_lm(model, typ=2)
            if doc:
                print_to_doc(doc, anova_table)
            else:
                print(anova_table)
        except Exception as e:
            error_msg = f"Error processing year {year}: {e}"
            if doc:
                print_to_doc(doc, error_msg)
            else:
                print(error_msg)

# Correlation Analysis
def correlation_analysis_with_relevant_columns(dataset, year, save_path='Charts/Correlation', doc=None):
    """Generates and saves correlation heatmaps for selected columns in the dataset."""
    if not os.path.exists(save_path):
        os.makedirs(save_path)

    relevant_columns = [
        'Total_AP', 'Total_PP', 'Total_SPA', 'Total_Bon',
        'Week1_AP', 'Week2_AP', 'Week3_AP', 'Week4_AP', 'Week5_AP',
        'Week1_PP', 'Week2_PP', 'Week3_PP', 'Week4_PP', 'Week5_PP'
    ]

    passed_students = dataset[dataset['Passed'] == 1]
    relevant_data = passed_students[relevant_columns + ['FinalGrade']]
    corr_matrix = relevant_data.corr()

    plt.figure(figsize=(12, 8))
    sns.heatmap(corr_matrix, annot=True, cmap='coolwarm', linewidths=0.5)
    plt.title(f'Correlation Matrix for {year} (Relevant Columns for Students that Passed)')
    image_path_passed = f'{save_path}/Correlation_Matrix_Passed_{year}.png'
    plt.savefig(image_path_passed, dpi=300)
    plt.show()

    if doc:
        add_image_to_doc(doc, image_path_passed)

    relevant_data = dataset[relevant_columns]
    corr_matrix = relevant_data.corr()

    plt.figure(figsize=(12, 8))
    sns.heatmap(corr_matrix, annot=True, cmap='coolwarm', linewidths=0.5)
    plt.title(f'Correlation Matrix for {year} (Relevant Columns - All Students)')
    image_path_all = f'{save_path}/Correlation_Matrix_{year}.png'
    plt.savefig(image_path_all, dpi=300)
    plt.show()

    if doc:
        add_image_to_doc(doc, image_path_all)

# Generate Weekly Heatmap
def generate_weekly_heatmap(corr_matrix, week, year, threshold, save_path='Charts/WeeklyHeatmap'):
    """Generates and saves a heatmap for weekly correlations above a given threshold."""
    if not os.path.exists(save_path):
        os.makedirs(save_path)

    filtered_corr_matrix = corr_matrix[(corr_matrix.abs() >= threshold).any(axis=1)]
    filtered_corr_matrix = filtered_corr_matrix.loc[:, (filtered_corr_matrix.abs() >= threshold).any(axis=0)]

    if filtered_corr_matrix.empty:
        print(f"No correlations above the threshold for {year} - Week {week}")
        return

    plt.figure(figsize=(12, 8))
    sns.heatmap(filtered_corr_matrix, annot=True, cmap='coolwarm', linewidths=0.8)
    plt.title(f'Correlation Matrix for {year} - Week {week} (Threshold: {threshold})')
    plt.savefig(f'{save_path}/Weekly_Correlation_Matrix_{year}_Week{week}.png', dpi=300)
    plt.show()

# Refined Correlation Analysis
def refined_correlation_analysis(dataset, year, weekly_lists, threshold=0.7, save_path='Charts/WeeklyHeatmap'):
    """Performs refined correlation analysis and generates heatmaps for each week."""
    for week, week_list in weekly_lists.items():
        print(f"\nGenerating heatmap for {year} - Week {week}")
        numerical_cols = dataset[week_list].select_dtypes(include=[np.number])
        corr_matrix = numerical_cols.corr()
        generate_weekly_heatmap(corr_matrix, week, year, threshold, save_path)

# Get Weekly Lists
def get_weekly_lists(all_lists, year):
    """Retrieves the weekly lists for the specified year."""
    try:
        weekly_lists = all_lists[str(year)]['weekly_and_participation_lists']['listWeekSync']
        return weekly_lists
    except KeyError as e:
        print(f"KeyError: {e} not found in all_lists for year {year}")
        return {}

# Feature Selection
def feature_selection(dataset, target_column, variance_threshold=0.8):
    """Selects top features based on variance threshold and mutual information."""
    exclude_columns = [
        'FinalGrade', 'FinalGradeInteger',
        'Total_AP', 'Total_PP', 'Total_SPA', 'Total_Bon',
        'Week1_AP', 'Week2_AP', 'Week3_AP', 'Week4_AP', 'Week5_AP',
        'Week1_PP', 'Week2_PP', 'Week3_PP', 'Week4_PP', 'Week5_PP'
    ]

    numerical_cols = dataset.select_dtypes(include=[np.number])
    numerical_cols = numerical_cols.drop(columns=exclude_columns, errors='ignore')
    numerical_cols = numerical_cols.fillna(numerical_cols.mean())

    sel = VarianceThreshold(threshold=(variance_threshold * (1 - variance_threshold)))
    numerical_cols_var = sel.fit_transform(numerical_cols)

    le = LabelEncoder()
    target_encoded = le.fit_transform(dataset[target_column])
    mi = mutual_info_classif(numerical_cols_var, target_encoded)

    feature_scores = pd.Series(mi, index=numerical_cols.columns[sel.get_support()]).sort_values(ascending=False)
    top_features = feature_scores.head(20).index.tolist()

    return top_features

# Hierarchical Clustering
def hierarchical_clustering_heatmap(dataset, year, features, save_path='Charts/HierarchicalClustering'):
    """Performs hierarchical clustering and generates a heatmap of the correlation matrix."""
    if not os.path.exists(save_path):
        os.makedirs(save_path)

    numerical_cols = dataset[features]
    corr_matrix = numerical_cols.corr()
    corr_matrix = corr_matrix.fillna(0)

    sns.clustermap(corr_matrix, annot=False, cmap='coolwarm', linewidths=0.5)
    plt.title(f'Hierarchical Clustering of Correlation Matrix for {year}')
    plt.savefig(f'{save_path}/Hierarchical_Clustering_{year}.png', dpi=300)
    plt.show()

# Plot Histograms
def plot_histogram(data, column, bins, title, xlabel, save_path='Charts/Histograms'):
    """Plots and saves a histogram for the specified column in the dataset."""
    if not os.path.exists(save_path):
        os.makedirs(save_path)

    if data[column].max() > 0:
        plt.figure(figsize=(12, 6))
        plt.hist(data[column], bins=bins, edgecolor='black', align='mid', color='skyblue')
        plt.title(f'{title} ({year})')
        plt.xlabel(xlabel)
        plt.ylabel('Frequency')
        plt.xticks(bins + 0.5, [str(int(b)) for b in bins + 0.5] if column != 'FinalGrade' else bins)
        plt.grid(False)
        plt.savefig(f'{save_path}/{title.replace(" ", "_")}_{year}.png', dpi=300)
        plt.show()

# Plot Categorical Attribute Distribution
def plot_categorical_attribute_distribution(datasets, column, save_path='Charts/CategoricalAttributes'):
    """Plots and saves the distribution of a categorical attribute across multiple years."""
    if not os.path.exists(save_path):
        os.makedirs(save_path)

    plt.figure(figsize=(14, 8))
    all_data = []

    for year, dataset in datasets.items():
        if column in dataset.columns:
            data = dataset[column].value_counts().reset_index()
            data.columns = [column, 'count']
            data['year'] = year
            all_data.append(data)

    combined_data = pd.concat(all_data)
    sns.barplot(x=column, y='count', hue='year', data=combined_data, palette='muted')
    plt.title(f'Distribution of {column} across years')
    plt.ylabel('Count')
    plt.xlabel(column)
    plt.xticks(rotation=45)
    plt.legend(title='Year')
    plt.tight_layout()
    plt.savefig(f'{save_path}/Distribution_{column}_across_years.png', dpi=300)
    plt.show()

# Custom function to replace 0/1 with Female/Male
def replace_sex_labels(df):
    """Replaces 0/1 in the 'Sex' column with 'Female'/'Male'."""
    df['Sex'] = df['Sex'].replace({'0': 'Female', '1': 'Male'})
    return df

# Average Points per Course
def average_points_per_course(datasets):
    """Calculates and prints the average points per course for each year."""
    for year, dataset in datasets.items():
        statistics = dataset.groupby('Course')[['Total_AP', 'Total_PP', 'Total_SPA', 'FinalGrade']].agg(['mean', 'min', 'max', 'std'])
        print(f"Year {year} - Points per course:\n{statistics}")

# Plot Grade Distribution
def plot_grade_distribution(dataset, year, save_path='Charts/Boxplots'):
    """Plots and saves the grade distribution by sex for each course."""
    dataset['Sex'] = dataset['Sex'].replace({'0': 'Female', '1': 'Male'})
    palette = {'Female': 'salmon', 'Male': '#89CFF0'}
    
    if not os.path.exists(save_path):
        os.makedirs(save_path)
    
    plt.figure(figsize=(10, 6))
    sns.boxplot(x='Course', y='FinalGrade', hue='Sex', data=dataset, palette=palette, linewidth=1, saturation=0.75)
    plt.title(f'Grade Distribution by Sex for {year}')
    plt.ylabel('Final Grades')
    plt.xlabel('Courses')
    plt.xticks(rotation=45)
    handles, labels = plt.gca().get_legend_handles_labels()
    plt.legend(handles, labels, title='Sex', loc='best')
    plt.tight_layout()
    
    plt.savefig(f'{save_path}/Grade_Distribution_{year}bySex_and_Course.png', dpi=300)
    plt.show()

# Plot Weekly Totals Histograms
def plot_weekly_totals(datasets, week_total, all_lists, save_path='Charts/WeeklyTotals'):
    """Plots and saves histograms for weekly totals for each year."""
    if not os.path.exists(save_path):
        os.makedirs(save_path)

    for year, dataset in datasets.items():
        weekly_list = all_lists[str(year)]['created_columns'][week_total]
        for col in weekly_list:
            if col in dataset.columns:
                max_value = int(dataset[col].max())
                min_value = int(dataset[col].min())
                if max_value > 0:
                    bins = np.arange(min_value, max_value + 2, max(1, (max_value - min_value) // 20)) - 0.5
                    plot_histogram(dataset, col, bins, f'Distribution of {col} for {year}', col, save_path)

# Plot Weekly Evolution
def plot_weekly_evolution(datasets, save_path='Charts/WeeklyEvolution'):
    """Plots and saves the weekly evolution of assessment points and participation points."""
    if not os.path.exists(save_path):
        os.makedirs(save_path)

    weeks = range(1, 6)
    metrics = ['AP', 'PP']
    years = sorted(datasets.keys())
    courses = datasets[next(iter(datasets))]['Course'].unique()

    for course in courses:
        for metric in metrics:
            mean_points = {week: [] for week in weeks}
            std_points = {week: [] for week in weeks}

            for week in weeks:
                for year in years:
                    dataset = datasets[year]
                    weekly_column = f'Week{week}_{metric}'

                    if weekly_column in dataset.columns:
                        course_data = dataset[dataset['Course'] == course]
                        weekly_mean = course_data[weekly_column].mean()
                        weekly_std = course_data[weekly_column].std()
                        mean_points[week].append(weekly_mean)
                        std_points[week].append(weekly_std)
                    else:
                        mean_points[week].append(np.nan)
                        std_points[week].append(np.nan)

            plt.figure(figsize=(14, 14))
            plt.subplot(2, 1, 1)
            for week in weeks:
                plt.bar(np.array(years) + (week - 1) * 0.1, mean_points[week], width=0.1, label=f'Week {week}')
            plt.title(f'Weekly Average {metric} Points for {course}')
            plt.xlabel('Year')
            plt.ylabel('Average Points')
            plt.xticks(years)
            plt.legend(title='Week', bbox_to_anchor=(1.05, 1), loc='upper left')
            plt.grid(True, axis='y')

            plt.subplot(2, 1, 2)
            for week in weeks:
                plt.bar(np.array(years) + (week - 1) * 0.1, std_points[week], width=0.1, label=f'Week {week}')
            plt.title(f'Weekly Standard Deviation of {metric} Points for {course}')
            plt.xlabel('Year')
            plt.ylabel('Standard Deviation')
            plt.xticks(years)
            plt.legend(title='Week', bbox_to_anchor=(1.05, 1), loc='upper left')
            plt.grid(True, axis='y')

            plt.tight_layout()
            plt.savefig(f'{save_path}/Weekly_Evolution_{course}_{metric}.png', dpi=300)
            plt.show()

# Plot Activity Types
def plot_activity_types(datasets, activity_type, all_lists, save_path='Charts/ActivityTypes'):
    """Plots and saves histograms for different activity types."""
    if not os.path.exists(save_path):
        os.makedirs(save_path)

    def plot_histogram(dataset, column, bins, title, xlabel, save_path):
        """Helper function to plot a histogram for a specified column."""
        plt.figure(figsize=(10, 6))
        sns.histplot(dataset[column].dropna(), bins=bins, kde=False)
        plt.title(title)
        plt.xlabel(xlabel)
        plt.ylabel('Frequency')
        plt.tight_layout()
        plt.savefig(f'{save_path}/{title.replace(" ", "_")}.png', dpi=300)
        plt.show()

    for year, dataset in datasets.items():
        activity_types = all_lists[str(year)]['activity_type_lists']
        if activity_type in activity_types:
            activity_list = activity_types[activity_type]
            for col in activity_list:
                if col in dataset.columns and dataset[col].max() > 0:
                    if activity_type in ['EvTypeA', 'EvTypeSPA']:
                        bins = np.arange(0, 3) - 0.5
                    elif activity_type == 'EvTypeB':
                        bins = np.arange(0, 5) - 0.5
                    elif activity_type == 'EvTypeC':
                        bins = np.arange(0, 7) - 0.5
                    else:
                        bins = np.arange(dataset[col].min(), dataset[col].max() + 1) - 0.5

                    plot_histogram(dataset, col, bins, f'Distribution of {col} for {year}', col, save_path)

# Plot Histograms with Custom Bins
def plot_histogram2(data, column, bins, title, xlabel, save_path='Charts/Histograms'):
    """Plots and saves a histogram for the specified column using custom bins."""
    if not os.path.exists(save_path):
        os.makedirs(save_path)

    if data[column].max() > 0:
        plt.figure(figsize=(12, 6))
        plt.hist(data[column], bins=bins, edgecolor='black', align='mid', color='skyblue')
        plt.title(title)
        plt.xlabel(xlabel)
        plt.ylabel('Frequency')
        plt.xticks(bins + 0.5, [str(int(b)) for b in bins + 0.5] if column != 'FinalGrade' else bins)
        plt.grid(False)
        plt.savefig(f'{save_path}/{title.replace(" ", "_")}.png', dpi=300)
        plt.show()

# Plot Final Grade Attributes
def plot_final_grade_attributes(datasets, save_path='Charts/FinalGrades'):
    """Plots and saves histograms for final grade attributes."""
    if not os.path.exists(save_path):
        os.makedirs(save_path)

    for year, dataset in datasets.items():
        passed = dataset[dataset['FinalGradeInteger'] > 0]
        if 'FinalGradeInteger' in passed.columns:
            bins = np.arange(10, 21) - 0.5
            plot_histogram2(passed, 'FinalGradeInteger', bins, f'Distribution of FinalGradeInteger for {year} (Passed Students)', 'FinalGradeInteger', save_path)
        if 'FinalGrade' in passed.columns:
            positive_final_grade = passed[passed['FinalGrade'] > 7.5]
            if len(positive_final_grade) > 8:
                bins = np.arange(7.5, 20.5, 0.5) - 0.5
                plot_histogram2(positive_final_grade, 'FinalGrade', bins, f'Distribution of FinalGrade for {year} (Positive Final Grades)', 'FinalGrade', save_path)
            else:
                print(f"Skipping {year} as there are less than 8 students with a positive final grade.")

# Function to plot and save the distribution of a categorical column
def plot_distribution(df, column, year, save_path='Charts/Distribution'):
    """Plots and saves the distribution of a categorical column as both a countplot and a pie chart."""
    if not os.path.exists(save_path):
        os.makedirs(save_path)

    plt.figure(figsize=(10, 6))
    sns.countplot(data=df, x=column, palette='deep')
    plt.title(f'Distribution of {column} for {year}')
    plt.xlabel(column)
    plt.ylabel('Count')
    plt.xticks(rotation=45, ha='right')
    plt.tight_layout()
    plt.savefig(f'{save_path}/Distribution_{column}_{year}.png', dpi=300)
    plt.show()

    plt.figure(figsize=(8, 8))
    value_counts = df[column].value_counts()
    plt.pie(value_counts, labels=value_counts.index, autopct='%1.1f%%', startangle=140, colors=sns.color_palette('deep'))
    plt.title(f'Distribution of {column} for {year}')
    plt.axis('equal')
    plt.tight_layout()
    plt.savefig(f'{save_path}/Pie_Distribution_{column}_{year}.png', dpi=300)
    plt.show()

# Function to plot histograms for a list of numerical columns
def plot_histograms(df, columns, save_path='Charts/Histograms'):
    """Plots and saves histograms for a list of numerical columns."""
    if not os.path.exists(save_path):
        os.makedirs(save_path)

    for column in columns:
        df[column].plot(kind='hist', bins=20, title=column, color='skyblue', edgecolor='black')
        plt.gca().spines[['top', 'right']].set_visible(False)
        plt.savefig(f'{save_path}/Histogram_{column}.png', dpi=300)
        plt.show()

# Function to calculate overall percentages
def calculate_overall_percentages(dataset):
    """Calculates the overall percentages for the 'Sex' column."""
    overall_percentages = dataset['Sex'].value_counts(normalize=True) * 100
    return overall_percentages

# Function to calculate fail percentages
def calculate_fail_percentages(dataset):
    """Calculates the fail percentages for the 'Sex' column."""
    failed_students = dataset[dataset['FinalGradeInteger'] < 10]
    fail_percentages = failed_students['Sex'].value_counts(normalize=True) * 100
    return fail_percentages

# Function to plot comparison by sex
def plot_comparison(dataset, year, course_name='CTCT', save_path='Charts/FailureRates'):
    """Plots and saves a comparison of overall and fail percentages by sex."""
    if not os.path.exists(save_path):
        os.makedirs(save_path)

    overall_percentages = calculate_overall_percentages(dataset)
    fail_percentages = calculate_fail_percentages(dataset)

    comparison = pd.DataFrame({
        'Overall_Percentage': overall_percentages,
        'Fail_Percentage': fail_percentages
    }).fillna(0)

    comparison.index = comparison.index.map({'0': 'Female', '1': 'Male'})

    comparison.plot(kind='bar', figsize=(12, 6), color=['#FF9999', '#89CFF0'])
    plt.title(f'Comparison of Percentages for {course_name} in {year} regarding Sex', fontsize=20)
    plt.xlabel('Sex', fontsize=16)
    plt.ylabel('Percentage', fontsize=16)
    plt.xticks(rotation=0)
    plt.tight_layout()
    plt.savefig(f'{save_path}/comparison_{course_name}_{year}.png', dpi=300)
    plt.show()

# Function to calculate and plot failure rates by course
def plot_failure_rates_by_course(dataset, year, save_path='Charts/FailureRates'):
    """Plots and saves the failure rates by course."""
    if not os.path.exists(save_path):
        os.makedirs(save_path)

    total_students_by_course = dataset['Course'].value_counts()
    failed_students_by_course = dataset[dataset['FinalGradeInteger'] < 10]['Course'].value_counts()

    overall_percentages_by_course = (total_students_by_course / total_students_by_course.sum()) * 100
    fail_percentages_by_course = (failed_students_by_course / failed_students_by_course.sum()) * 100

    comparison = pd.DataFrame({
        'Overall_Percentage': overall_percentages_by_course,
        'Fail_Percentage': fail_percentages_by_course
    }).fillna(0)

    comparison.plot(kind='bar', figsize=(12, 6), color=['#FF9999', '#89CFF0'])
    plt.title(f'Comparison of Failure Rates by Course in {year}', fontsize=20)
    plt.xlabel('Course', fontsize=16)
    plt.ylabel('Percentage', fontsize=16)
    plt.xticks(rotation=45)
    plt.tight_layout()
    plt.savefig(f'{save_path}/failure_rates_by_course_{year}.png', dpi=300)
    plt.show()

# Function to perform regression analysis and generate related plots
def regression_analysis(df, x_col, y_col, year, save_path='Charts/Regression', doc=None):
    """Performs regression analysis and generates related plots."""
    if not os.path.exists(save_path):
        os.makedirs(save_path)

    df = df[[x_col, y_col]].dropna()

    if df.empty:
        error_msg = f"No data available for regression analysis between {x_col} and {y_col} for year {year}."
        if doc:
            print_to_doc(doc, error_msg)
        else:
            print(error_msg)
        return

    X = df[x_col].values.reshape(-1, 1)
    y = df[y_col].values.reshape(-1, 1)
    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=0)
    regressor = LinearRegression()
    regressor.fit(X_train, y_train)
    y_pred = regressor.predict(X_test)

    plt.scatter(X_train, y_train, color='red')
    plt.plot(X_train, regressor.predict(X_train), color='blue')
    plt.title(f'{x_col} vs {y_col} (Training Set) ({year})')
    plt.xlabel(x_col)
    plt.ylabel(y_col)
    image_path_train = f'{save_path}/{x_col}_vs_{y_col}_Training_{year}.png'
    plt.savefig(image_path_train, dpi=300)
    plt.show()

    if doc:
        add_image_to_doc(doc, image_path_train)

    plt.scatter(X_test, y_test, color='red')
    plt.plot(X_test, y_pred, color='blue')
    plt.title(f'{x_col} vs {y_col} (Test Set) ({year})')
    plt.xlabel(x_col)
    plt.ylabel(y_col)
    image_path_test = f'{save_path}/{x_col}_vs_{y_col}_Test_{year}.png'
    plt.savefig(image_path_test, dpi=300)
    plt.show()

    if doc:
        add_image_to_doc(doc, image_path_test)

    r_squared = regressor.score(X_test, y_test)
    intercept = regressor.intercept_[0]
    slope = regressor.coef_[0][0]
    result = (
        f"Regression Analysis for {year}:\n"
        f"{x_col} vs {y_col}:\n"
        f"R-squared: {r_squared}\n"
        f"Intercept: {intercept}\n"
        f"Slope: {slope}\n"
    )
    if doc:
        print_to_doc(doc, result)
    else:
        print(result)

# Function to perform logistic regression analysis
def logistic_regression_analysis(df, x_cols, y_col, year, threshold=0.6, test_size=0.2, save_path='Charts/LogisticRegression', doc=None):
    """Performs logistic regression analysis and generates related plots."""
    if not os.path.exists(save_path):
        os.makedirs(save_path)

    df = df[x_cols + [y_col]].dropna()

    if df.empty:
        error_msg = f"No data available for logistic regression analysis between {x_cols} and {y_col} for year {year}."
        if doc:
            print_to_doc(doc, error_msg)
        else:
            print(error_msg)
        return

    X = df[x_cols].values
    y = df[y_col].values
    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=test_size, random_state=0)

    classifier = LogisticRegression()
    classifier.fit(X_train, y_train)
    y_pred = classifier.predict(X_test)

    accuracy = accuracy_score(y_test, y_pred)

    if accuracy >= threshold:
        conf_matrix = confusion_matrix(y_test, y_pred)
        plt.figure(figsize=(8, 6))
        plt.imshow(conf_matrix, cmap='Blues', interpolation='nearest')
        plt.title(f'Confusion Matrix ({year})')
        plt.xlabel('Predicted Label')
        plt.ylabel('True Label')
        plt.colorbar()
        for i in range(len(conf_matrix)):
            for j in range(len(conf_matrix[i])):
                plt.text(j, i, conf_matrix[i][j], ha='center', va='center', color='red')
        image_path_conf_matrix = f'{save_path}/Confusion_Matrix_{year}.png'
        plt.savefig(image_path_conf_matrix, dpi=300)
        plt.show()

        if doc:
            add_image_to_doc(doc, image_path_conf_matrix)

        class_report = classification_report(y_test, y_pred)
        result = (
            f"Logistic Regression Analysis for {year}:\n"
            f"{x_cols} vs {y_col}:\n"
            f"Accuracy: {accuracy}\n"
            f"Classification Report:\n{class_report}\n"
        )
        if doc:
            print_to_doc(doc, result)
        else:
            print(result)
    else:
        no_result_msg = f"Accuracy below threshold for {x_cols} vs {y_col} in year {year}. Skipping documentation."
        if doc:
            print_to_doc(doc, no_result_msg)
        else:
            print(no_result_msg)

# Function to perform random forest analysis
def random_forest_analysis(df, x_cols, y_col, year, save_path='Charts/RandomForest', doc=None):
    """Performs random forest analysis and generates related plots."""
    if not os.path.exists(save_path):
        os.makedirs(save_path)

    df = df[x_cols + [y_col]].dropna()

    if df.empty:
        error_msg = f"No data available for random forest analysis between {x_cols} and {y_col} for year {year}."
        if doc:
            print_to_doc(doc, error_msg)
        else:
            print(error_msg)
        return

    X = df[x_cols].values
    y = df[y_col].values
    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=0)

    classifier = RandomForestClassifier(n_estimators=100, random_state=0)
    classifier.fit(X_train, y_train)
    y_pred = classifier.predict(X_test)

    accuracy = accuracy_score(y_test, y_pred)

    conf_matrix = confusion_matrix(y_test, y_pred)
    plt.figure(figsize=(8, 6))
    sns.heatmap(conf_matrix, annot=True, cmap='Blues', fmt='g')
    plt.title(f'Confusion Matrix ({year})')
    plt.xlabel('Predicted')
    plt.ylabel('True')
    image_path_conf_matrix = f'{save_path}/Confusion_Matrix_{year}.png'
    plt.savefig(image_path_conf_matrix, dpi=300)
    plt.show()

    if doc:
        add_image_to_doc(doc, image_path_conf_matrix)

    class_report = classification_report(y_test, y_pred)
    result = (
        f"Random Forest Analysis for {year}:\n"
        f"{x_cols} vs {y_col}:\n"
        f"Accuracy: {accuracy}\n"
        f"Classification Report:\n{class_report}\n"
    )
    if doc:
        print_to_doc(doc, result)
    else:
        print(result)

# Enhanced Clustering Analysis Functions
def save_and_show_plot(fig, image_name, save_path='Charts'):
    if not os.path.exists(save_path):
        os.makedirs(save_path)
    fig.savefig(os.path.join(save_path, image_name))
    plt.close(fig)

# Function to perform feature selection
def select_features(dataset, target_column='FinalGrade', variance_threshold=0.8, top_n=20):
    """Selects top features based on variance threshold and mutual information."""
    exclude_columns = ['FinalGrade', 'FinalGradeInteger', 'Sex', 'Course']
    numerical_cols = dataset.select_dtypes(include=[np.number])
    numerical_cols = numerical_cols.drop(columns=exclude_columns, errors='ignore')
    numerical_cols = numerical_cols.fillna(numerical_cols.mean())
    
    sel = VarianceThreshold(threshold=(variance_threshold * (1 - variance_threshold)))
    numerical_cols_var = sel.fit_transform(numerical_cols)
    
    le = LabelEncoder()
    target_encoded = le.fit_transform(dataset[target_column])
    mi = mutual_info_classif(numerical_cols_var, target_encoded)
    
    feature_scores = pd.Series(mi, index=numerical_cols.columns[sel.get_support()]).sort_values(ascending=False)
    top_features = feature_scores.head(top_n).index.tolist()
    
    return top_features

def clustering_analysis(df, features, year, save_path='Charts/Clustering'):
    """Performs clustering analysis and generates related plots."""
    if not os.path.exists(save_path):
        os.makedirs(save_path)

    X = df[features].values

    wcss = []
    for i in range(1, 11):
        kmeans = KMeans(n_clusters=i, init='k-means++', random_state=42, n_init=10)
        kmeans.fit(X)
        wcss.append(kmeans.inertia_)

    plt.figure()
    plt.plot(range(1, 11), wcss)
    plt.title(f'Elbow Method for Determining Optimal Clusters ({year})')
    plt.xlabel('Number of clusters')
    plt.ylabel('WCSS')
    plt.savefig(f'{save_path}/Elbow_Method_{year}.png')
    plt.show()

    kmeans = KMeans(n_init=10, n_clusters=2, init='k-means++', random_state=0)
    y_kmeans = kmeans.fit_predict(X)
    plt.scatter(X[y_kmeans == 0, 0], X[y_kmeans == 0, 1], s=30, c='red', label='Cluster 1')
    plt.scatter(X[y_kmeans == 1, 0], X[y_kmeans == 1, 1], s=20, c='blue', label='Cluster 2')
    plt.scatter(kmeans.cluster_centers_[:, 0], kmeans.cluster_centers_[:, 1], s=100, c='yellow', label='Centroids')
    plt.title(f'K-Means Clustering of Students ({year})')
    plt.xlabel(features[0])
    plt.ylabel(features[1])
    plt.legend()
    plt.savefig(f'{save_path}/KMeans_Clusters_{year}.png')
    plt.show()

    plt.figure()
    dendrogram = sch.dendrogram(sch.linkage(X, method='ward'))
    plt.title(f'Dendrogram ({year})')
    plt.xlabel('Students')
    plt.ylabel('Euclidean distances')
    plt.savefig(f'{save_path}/Dendrogram_{year}.png')
    plt.show()

    hc = AgglomerativeClustering(n_clusters=2, metric='euclidean', linkage='ward')
    y_hc = hc.fit_predict(X)
    plt.scatter(X[y_hc == 0, 0], X[y_hc == 0, 1], s=30, c='green', label='Cluster 1')
    plt.scatter(X[y_hc == 1, 0], X[y_hc == 1, 1], s=20, c='purple', label='Cluster 2')
    plt.title(f'Hierarchical Clustering of Students ({year})')
    plt.xlabel(features[0])
    plt.ylabel(features[1])
    plt.legend()
    plt.savefig(f'{save_path}/Hierarchical_Clusters_{year}.png')
    plt.show()

import pandas as pd
import os
import numpy as np
import matplotlib.pyplot as plt
from sklearn.cluster import KMeans, AgglomerativeClustering
from sklearn.preprocessing import StandardScaler, LabelEncoder
from sklearn.decomposition import PCA
from sklearn.metrics import silhouette_score, davies_bouldin_score, calinski_harabasz_score
from scipy.cluster.hierarchy import dendrogram, linkage
from sklearn.feature_selection import VarianceThreshold, mutual_info_classif

# Enhanced Clustering Analysis Functions
def save_and_show_plot(fig, image_name, save_path='Charts'):
    if not os.path.exists(save_path):
        os.makedirs(save_path)
    fig.savefig(os.path.join(save_path, image_name))
    plt.close(fig)

# Function to perform feature selection
def select_features(dataset, target_column='FinalGrade', variance_threshold=0.8, top_n=20):
    """Selects top features based on variance threshold and mutual information."""
    exclude_columns = ['FinalGrade', 'FinalGradeInteger', 'Sex', 'Course']
    numerical_cols = dataset.select_dtypes(include=[np.number])
    numerical_cols = numerical_cols.drop(columns=exclude_columns, errors='ignore')
    numerical_cols = numerical_cols.fillna(numerical_cols.mean())
    
    sel = VarianceThreshold(threshold=(variance_threshold * (1 - variance_threshold)))
    numerical_cols_var = sel.fit_transform(numerical_cols)
    
    le = LabelEncoder()
    target_encoded = le.fit_transform(dataset[target_column])
    mi = mutual_info_classif(numerical_cols_var, target_encoded)
    
    feature_scores = pd.Series(mi, index=numerical_cols.columns[sel.get_support()]).sort_values(ascending=False)
    top_features = feature_scores.head(top_n).index.tolist()
    
    return top_features

# Function to display PCA components and their contribution to the Principal Components
def display_pca_components(pca, feature_names):
    components_df = pd.DataFrame(pca.components_, columns=feature_names, index=[f'PC{i+1}' for i in range(pca.n_components_)])
    explained_variance = pca.explained_variance_ratio_
    print("PCA Components and their contribution to the Principal Components:")
    print(components_df)
    print("\nExplained variance ratio per Principal Component:")
    for i, variance in enumerate(explained_variance):
        print(f"PC{i+1}: {variance:.4f}")

# Function to plot the Elbow Method for determining the optimal number of clusters
def plot_elbow_method(X, year, max_clusters=10):
    """Plots the Elbow Method for determining the optimal number of clusters."""
    wcss = []
    for i in range(1, max_clusters + 1):
        kmeans = KMeans(n_clusters=i, init='k-means++', random_state=42, n_init=10)
        kmeans.fit(X)
        wcss.append(kmeans.inertia_)

    fig, ax = plt.subplots()
    ax.plot(range(1, max_clusters + 1), wcss, marker='o')
    ax.set_title(f'Elbow Method for Determining Optimal Clusters ({year})')
    ax.set_xlabel('Number of clusters')
    ax.set_ylabel('WCSS')

    image_name = f'Elbow_Method_{year}.png'
    save_and_show_plot(fig, image_name)

# Function to perform K-Means clustering and plot the results
def kmeans_clustering(X, features, year, n_clusters=2):
    """Performs K-Means clustering and generates related plots."""
    kmeans = KMeans(n_clusters=n_clusters, init='k-means++', random_state=0, n_init=10)
    y_kmeans = kmeans.fit_predict(X)

    silhouette_avg = silhouette_score(X, y_kmeans)
    davies_bouldin = davies_bouldin_score(X, y_kmeans)
    calinski_harabasz = calinski_harabasz_score(X, y_kmeans)

    print(f"Year: {year} | K-Means Clustering | Clusters: {n_clusters}")
    print(f"Silhouette Score: {silhouette_avg}")
    print(f"Davies-Bouldin Index: {davies_bouldin}")
    print(f"Calinski-Harabasz Index: {calinski_harabasz}")

    fig, ax = plt.subplots()
    for cluster in range(n_clusters):
        ax.scatter(X[y_kmeans == cluster, 0], X[y_kmeans == cluster, 1], s=30, label=f'Cluster {cluster + 1}')
    ax.scatter(kmeans.cluster_centers_[:, 0], kmeans.cluster_centers_[:, 1], s=100, c='yellow', label='Centroids')
    ax.set_title(f'K-Means Clustering of Students ({year})')
    ax.set_xlabel('Principal Component 1')
    ax.set_ylabel('Principal Component 2')
    ax.legend()

    image_name = f'KMeans_Clusters_{n_clusters}_{year}.png'
    save_and_show_plot(fig, image_name)

# Function to plot Hierarchical Clustering Dendrogram
def plot_dendrogram(X, year):
    """Plots a dendrogram for hierarchical clustering."""
    Z = linkage(X, method='ward')
    fig, ax = plt.subplots()
    dendrogram(Z, ax=ax)
    ax.set_title(f'Dendrogram ({year})')
    ax.set_xlabel('Students')
    ax.set_ylabel('Euclidean distances')

    image_name = f'Dendrogram_{year}.png'
    save_and_show_plot(fig, image_name)

# Function to perform Hierarchical Clustering and plot the results
def hierarchical_clustering(X, features, year, n_clusters=2):
    """Performs hierarchical clustering and generates related plots."""
    hc = AgglomerativeClustering(n_clusters=n_clusters, metric='euclidean', linkage='ward')
    y_hc = hc.fit_predict(X)

    silhouette_avg = silhouette_score(X, y_hc)
    davies_bouldin = davies_bouldin_score(X, y_hc)
    calinski_harabasz = calinski_harabasz_score(X, y_hc)

    print(f"Year: {year} | Hierarchical Clustering | Clusters: {n_clusters}")
    print(f"Silhouette Score: {silhouette_avg}")
    print(f"Davies-Bouldin Index: {davies_bouldin}")
    print(f"Calinski-Harabasz Index: {calinski_harabasz}")

    fig, ax = plt.subplots()
    for cluster in range(n_clusters):
        ax.scatter(X[y_hc == cluster, 0], X[y_hc == cluster, 1], s=30, label=f'Cluster {cluster + 1}')
    ax.set_title(f'Hierarchical Clustering of Students ({year})')
    ax.set_xlabel('Principal Component 1')
    ax.set_ylabel('Principal Component 2')
    ax.legend()

    image_name = f'Hierarchical_Clusters_{n_clusters}_{year}.png'
    save_and_show_plot(fig, image_name)

# Function to perform extensive clustering analysis
def extensive_clustering_analysis(df, features, year):
    """Performs extensive clustering analysis and generates related plots."""
    existing_features = [feature for feature in features if feature in df.columns]
    if not existing_features:
        print(f"No valid features for clustering in year {year}")
        return

    if len(existing_features) < 2:
        print(f"Not enough features for clustering in year {year}")
        return

    scaler = StandardScaler()
    X = scaler.fit_transform(df[existing_features])

    if X.shape[0] < 2:
        print(f"Not enough samples for clustering in year {year}")
        return

    pca = PCA(n_components=2)
    X_pca = pca.fit_transform(X)
    display_pca_components(pca, existing_features)

    plot_elbow_method(X, year)
    kmeans_clustering(X_pca, existing_features, year, n_clusters=2)
    plot_dendrogram(X, year)
    hierarchical_clustering(X_pca, existing_features, year, n_clusters=2)

# Encoding categorical variables
def encode_categorical_variables(df, columns):
    """Encodes categorical variables as numerical codes."""
    for col in columns:
        df[col] = df[col].astype('category').cat.codes
    return df

# Discretizing continuous variables
def discretize_continuous_variables(df, columns, bins=10):
    """Discretizes continuous variables into specified number of bins."""
    for col in columns:
        df[col] = pd.cut(df[col], bins, labels=False)
    return df

# Scaling features
def scale_features(df, columns):
    """Scales continuous features using StandardScaler."""
    scaler = StandardScaler()
    df[columns] = scaler.fit_transform(df[columns])
    return df
